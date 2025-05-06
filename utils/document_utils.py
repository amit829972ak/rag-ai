import pandas as pd
import io
import json
import os
from typing import Dict, List, Union, Tuple, Optional
import csv
import PyPDF2
from docx import Document
import base64
from tabulate import tabulate

def process_document(file, file_type: str) -> Tuple[str, Optional[pd.DataFrame]]:
    """
    Process various document types and extract their content.
    
    Args:
        file: The uploaded file object
        file_type (str): The file extension/type
        
    Returns:
        Tuple[str, Optional[pd.DataFrame]]: Tuple containing extracted text content and dataframe (if applicable)
    """
    content = ""
    df = None
    
    try:
        # Handle different file types
        if file_type in ['.csv', '.tsv']:
            # CSV/TSV files
            delimiter = ',' if file_type == '.csv' else '\t'
            df = pd.read_csv(file, delimiter=delimiter)
            # Create a string representation
            content = f"Document type: {file_type[1:].upper()}\n\n"
            content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
            content += f"Column names: {', '.join(df.columns.tolist())}\n\n"
            # Add full tabular content
            content += tabulate(df, headers='keys', tablefmt='pipe', showindex=False)
        
        elif file_type == '.xlsx':
            # Excel files
            # Read all sheets
            xlsx = pd.ExcelFile(file, engine='openpyxl')
            sheet_names = xlsx.sheet_names
            content = f"Document type: Excel\n\n"
            content += f"Total sheets: {len(sheet_names)}\n"
            content += f"Sheet names: {', '.join(sheet_names)}\n\n"
            
            # Process each sheet
            for sheet_name in sheet_names:
                df_sheet = pd.read_excel(file, engine='openpyxl', sheet_name=sheet_name)
                content += f"=== SHEET: {sheet_name} ===\n"
                content += f"Rows: {len(df_sheet)}, Columns: {len(df_sheet.columns)}\n"
                content += f"Column names: {', '.join(df_sheet.columns.tolist())}\n\n"
                # Add full tabular content
                content += tabulate(df_sheet, headers='keys', tablefmt='pipe', showindex=False)
                content += "\n\n"
                
            # Use the first sheet as the default DataFrame
            df = pd.read_excel(file, engine='openpyxl', sheet_name=0)
            
        elif file_type == '.pdf':
            # PDF files
            reader = PyPDF2.PdfReader(file)
            content = f"Document type: PDF\n\n"
            content += f"Total pages: {len(reader.pages)}\n\n"
            
            # Extract text from all pages
            for i in range(len(reader.pages)):
                page_content = reader.pages[i].extract_text()
                content += f"--- Page {i+1} ---\n{page_content}\n\n"
                
        elif file_type == '.txt':
            # Text files
            content = f"Document type: Text file\n\n"
            text_content = file.read().decode('utf-8')
            # Include the full content without truncation
            content += text_content
                
        elif file_type == '.docx':
            # Word documents
            doc = Document(file)
            content = f"Document type: Word Document\n\n"
            
            # Extract all paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content += "\n\n".join(paragraphs)
            
            # Include all tables if present
            if doc.tables:
                content += "\n\n=== TABLES ===\n\n"
                for i, table in enumerate(doc.tables):
                    content += f"--- Table {i+1} ---\n"
                    for row in table.rows:
                        row_content = [cell.text for cell in row.cells]
                        content += " | ".join(row_content) + "\n"
                    content += "\n"
                
        else:
            content = f"Unsupported file type: {file_type}"
            
        return content, df
    
    except Exception as e:
        return f"Error processing {file_type} file: {str(e)}", None

def get_file_extension(filename: str) -> str:
    """
    Get the lowercase file extension including the dot.
    
    Args:
        filename (str): The filename
        
    Returns:
        str: Lowercase file extension with dot
    """
    _, extension = os.path.splitext(filename)
    return extension.lower()

def convert_df_to_csv_string(df: pd.DataFrame) -> str:
    """
    Convert a DataFrame to a CSV string.
    
    Args:
        df (pd.DataFrame): DataFrame to convert
        
    Returns:
        str: CSV string representation
    """
    if df is None:
        return ""
    
    csv_buffer = io.StringIO()
    df.to_csv(csv_buffer, index=False)
    return csv_buffer.getvalue()

def convert_df_to_json_string(df: pd.DataFrame) -> str:
    """
    Convert a DataFrame to a JSON string.
    
    Args:
        df (pd.DataFrame): DataFrame to convert
        
    Returns:
        str: JSON string representation
    """
    if df is None:
        return ""
    
    return df.to_json(orient='records')

def get_document_summary(document_content: str) -> str:
    """
    Create a brief summary of the document content.
    
    Args:
        document_content (str): The document content
        
    Returns:
        str: A brief summary
    """
    # Extract document type from content
    doc_type = "Unknown"
    if document_content.startswith("Document type:"):
        doc_type_line = document_content.split('\n')[0]
        doc_type = doc_type_line.replace("Document type:", "").strip()
    
    # Get a length-based summary
    content_length = len(document_content)
    words = document_content.split()
    word_count = len(words)
    
    # Create summary
    summary = f"{doc_type} document with approximately {word_count} words"
    
    # Add additional info based on document type
    if "Rows:" in document_content and "Columns:" in document_content:
        # For tabular data
        info_line = [line for line in document_content.split('\n') if "Rows:" in line][0]
        summary += f" ({info_line.strip()})"
    elif "Total pages:" in document_content:
        # For PDFs
        info_line = [line for line in document_content.split('\n') if "Total pages:" in line][0]
        summary += f" ({info_line.strip()})"
    
    return summary
